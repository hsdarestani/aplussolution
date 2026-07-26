from __future__ import annotations

import hashlib
import json
import os
import tempfile
import urllib.request
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase.acroform import AcroForm
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

OUT = Path(os.environ.get('TEMPLATE_BUNDLE_OUTPUT', '/tmp/aplus-document-templates.zip'))
WORK = Path(tempfile.mkdtemp(prefix='aplus-private-templates-'))


def sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def base_doc(title: str) -> Document:
    doc = Document()
    styles = doc.styles
    styles['Normal'].font.name = 'DejaVu Sans'
    styles['Normal'].font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)
    return doc


def add_signature(doc: Document, roles=('Arbeitnehmer/in', 'Arbeitgeber')):
    doc.add_paragraph('')
    doc.add_paragraph('Ort, Datum: {{ signature_place }}, {{ signature_date }}')
    for role in roles:
        doc.add_paragraph(f'Unterschrift {role}: ______________________________')


def save_doc(doc: Document, filename: str) -> Path:
    path = WORK / filename
    doc.save(path)
    return path


def make_employment_contract() -> Path:
    d = base_doc('Arbeitsvertrag mit Inbezugnahme des DGB/GVP-Tarifwerks')
    d.add_paragraph('Stand: November 2025 – zur Verwendung durch A+ Solution GmbH')
    d.add_paragraph('zwischen')
    d.add_paragraph('Firma: {{ company_name }}')
    d.add_paragraph('Anschrift: {{ company_address }}')
    d.add_paragraph('– im Folgenden „Arbeitgeber“ genannt –')
    d.add_paragraph('und')
    d.add_paragraph('Name: {{ employee_name }}')
    d.add_paragraph('Anschrift: {{ employee_address }}')
    d.add_paragraph('– im Folgenden „Mitarbeiter“ genannt –')
    sections = [
        ('1. Vertragsgegenstand', [
            'Der Arbeitgeber überlässt seinen Kunden Personal für die Durchführung von Arbeiten im Rahmen der Arbeitnehmerüberlassung. Der Arbeitgeber besitzt eine Erlaubnis zur Arbeitnehmerüberlassung, zuletzt erteilt von {{ aueg_license_authority }} am {{ aueg_license_date }}.',
            'Der Mitarbeiter wird in der Regel auswärts an wechselnden Einsatzstellen in Kundenbetrieben und bei wechselnden Kunden eingesetzt. Vereinbarter Einsatzbereich: {{ einsatzbereich }}.',
            'Der Arbeitgeber ist berechtigt, den Mitarbeiter von seinem Einsatzort abzuberufen und ihn anderweitig einzusetzen.',
        ]),
        ('2. Inbezugnahme von Tarifverträgen / Kollisionsklauseln', [
            'Die Rechte und Pflichten bestimmen sich nach den jeweils gültigen Tarifverträgen, die der Gesamtverband der Personaldienstleister e.V. (GVP) mit den Gewerkschaften der DGB-Tarifgemeinschaft Zeitarbeit abgeschlossen hat, sowie den ergänzenden, ändernden oder ersetzenden Tarifverträgen.',
            'Ab Beginn eines Einsatzes gelten die Tarifverträge der nach ihrer Satzung für den Einsatzbetrieb zuständigen Gewerkschaft. Bis zum ersten Einsatz finden die Tarifverträge unter Beteiligung von ver.di Anwendung. Günstigere einzelvertragliche Regelungen bleiben unberührt.',
        ]),
        ('3. Beginn und Beendigung', [
            'Das Arbeitsverhältnis beginnt am {{ start_date }} und ist – soweit ein Enddatum vereinbart ist – befristet bis {{ end_date }}.',
            'Neueinstellung: {{ neuanstellung }}. Die Kündigungsfristen und das Kündigungsverfahren richten sich nach den anwendbaren Tarifverträgen und gesetzlichen Vorschriften. Kündigungen bedürfen der Schriftform.',
        ]),
        ('4. Tätigkeit und Eingruppierung', [
            'Tätigkeit(en): {{ taetigkeit }}.',
            'Eingruppierung nach den anwendbaren Entgeltrahmentarifverträgen in Entgeltgruppe {{ pay_group }}.',
        ]),
        ('5. Arbeitszeit', [
            'Beschäftigungsart: {{ employment_type }}. Individuelle regelmäßige monatliche Arbeitszeit: {{ monthly_hours }} Stunden.',
            'Ruhepausen, Ruhezeiten, Schichtsystem, Schichtrhythmus, Schichtänderungen sowie Überstunden und Mehrarbeit richten sich nach den anwendbaren Manteltarifverträgen.',
            'Bei geringfügiger Beschäftigung können Entgelterhöhungen durch Anpassung der monatlichen Arbeitszeit ausgeglichen werden, soweit dies zur Einhaltung der gesetzlichen Entgeltgrenze erforderlich ist.',
        ]),
        ('6. Vergütung und Fälligkeit', [
            'Tarifliches Stundenentgelt: {{ tariff_hourly_rate }} Euro. Die Auszahlung erfolgt unbar auf das vom Mitarbeiter angegebene Konto.',
        ]),
        ('7. Übertarifliche Zulage', [
            'Übertarifliche Zulage je tatsächlich geleisteter Arbeitsstunde: {{ extra_allowance }} Euro. Die Anrechenbarkeit und ein möglicher Widerruf richten sich nach den vereinbarten und tariflichen Bestimmungen.',
        ]),
        ('8. Urlaub', ['Dauer und Berechnung des Erholungsurlaubs richten sich nach den jeweils anwendbaren Manteltarifverträgen.']),
        ('9. Nebentätigkeit', ['Jede Nebentätigkeit ist anzuzeigen. Sie kann untersagt werden, wenn arbeitsvertragliche Leistungen oder betriebliche Interessen beeinträchtigt werden.']),
        ('10. Vertragsstrafe', ['Löst der Mitarbeiter das Arbeitsverhältnis schuldhaft ohne Rechtsgrund und ohne Einhaltung der Kündigungsfrist, gelten die im anwendbaren Tarifwerk und Arbeitsvertrag vorgesehenen Rechtsfolgen.']),
        ('11. Ausschlussfristen', ['Ansprüche aus dem Arbeitsverhältnis sind innerhalb der tariflich bzw. arbeitsvertraglich vorgesehenen Fristen geltend zu machen. Gesetzlich zwingende Ansprüche bleiben unberührt.']),
        ('12. Sonstiges', ['Wesentliche Änderungen persönlicher Verhältnisse sind unverzüglich in Textform mitzuteilen. Nebenabreden und Änderungen bedürfen mindestens der Textform; Befristungsabreden der gesetzlich erforderlichen Form.']),
        ('13. Zusätzliche Vereinbarungen', ['{{ additional_agreements }}']),
        ('14. Zusätzliche Dokumente', ['Als Anlage erhält der Mitarbeiter das Merkblatt der Erlaubnisbehörde über den wesentlichen Inhalt des AÜG.']),
    ]
    for heading, paragraphs in sections:
        p = d.add_paragraph()
        r = p.add_run(heading)
        r.bold = True
        for text in paragraphs:
            d.add_paragraph(text)
    add_signature(d)
    return save_doc(d, 'AV Muster 20262027.docx')


def make_termination() -> Path:
    d = base_doc('Aufhebungsvertrag')
    d.add_paragraph('Zwischen {{ company_name }}, {{ company_address }} – Arbeitgeber – und {{ employee_name }}, {{ employee_address }} – Arbeitnehmer – wird folgender Aufhebungsvertrag geschlossen:')
    for heading, text in [
        ('§ 1 Beendigung des Arbeitsverhältnisses', 'Das bestehende Arbeitsverhältnis wird im gegenseitigen Einvernehmen zum {{ termination_date }} beendet. Die vereinbarte Kündigungsfrist wurde berücksichtigt.'),
        ('§ 2 Arbeitsfreistellung', 'Arbeitsfreistellung bis Vertragsende: {{ release_from_work }}. Soweit vereinbart, erfolgt sie unter Fortzahlung der vertraglichen Vergütung und unter Anrechnung bestehender Freistellungs- und Urlaubsansprüche.'),
        ('§ 3 Urlaub', 'Noch bestehender Resturlaub: {{ remaining_leave_days }} Tag(e). Dieser wird nach Maßgabe der Vereinbarung während der Freistellung gewährt.'),
        ('§ 4 Wettbewerbsvereinbarung', 'Bestehende wirksame Wettbewerbsvereinbarungen bleiben unberührt.'),
        ('§ 5 Ausgleich aller Ansprüche', 'Mit Erfüllung dieser Vereinbarung sind die wechselseitigen finanziellen Ansprüche aus dem Arbeitsverhältnis und seiner Beendigung abgegolten, soweit dies rechtlich zulässig ist.'),
    ]:
        p = d.add_paragraph(); p.add_run(heading).bold = True; d.add_paragraph(text)
    add_signature(d)
    return save_doc(d, 'Aufhebungsvertrag - Muster.docx')


def make_addendum() -> Path:
    d = base_doc('Zusatzvereinbarung – Vergütung nach tatsächlich geleisteten Stunden (tarifkonform)')
    d.add_paragraph('Zwischen {{ company_name }}, {{ company_address }} – Arbeitgeber – und {{ employee_name }} – Arbeitnehmer –')
    clauses = [
        ('§ 1 Tarifliche Grundlage', 'Auf das Arbeitsverhältnis finden die Tarifverträge DGB/GVP in ihrer jeweils gültigen Fassung Anwendung.'),
        ('§ 2 Regelmäßige Arbeitszeit', 'Die individuell vereinbarte regelmäßige monatliche Arbeitszeit beträgt {{ monthly_hours }} Stunden.'),
        ('§ 3 Abrechnung der Vergütung', 'Die Vergütung wird auf Grundlage der tatsächlich geleisteten und ordnungsgemäß erfassten Arbeitsstunden abgerechnet. Zum Ausgleich von Abweichungen wird ein Arbeitszeitkonto gemäß § 4 MTV DGB/GVP geführt.'),
        ('§ 4 Voraussetzungen des Vergütungsanspruchs', 'Ein Vergütungsanspruch besteht, wenn der Arbeitnehmer arbeitsfähig und arbeitsbereit ist und einen zumutbaren Einsatz nicht ohne sachlichen Grund ablehnt.'),
        ('§ 5 Annahmeverzug', 'Zeiten, in denen der Arbeitnehmer seine Arbeitsleistung ordnungsgemäß anbietet, jedoch kein Einsatz erfolgt, werden gemäß § 11 Abs. 4 AÜG vergütet.'),
        ('§ 6 Schlussbestimmung', 'Im Übrigen gelten der Arbeitsvertrag sowie die Tarifverträge DGB/GVP.'),
    ]
    for h, t in clauses:
        p = d.add_paragraph(); p.add_run(h).bold = True; d.add_paragraph(t)
    add_signature(d)
    return save_doc(d, '2. Zusatzvereinbarung - Entgelt .docx')


def make_rehire() -> Path:
    d = base_doc('Vereinbarung zur Wiederaufnahme eines beendeten Arbeitsverhältnisses')
    d.add_paragraph('Zwischen {{ company_name }}, {{ company_address }} – Arbeitgeber – und {{ employee_name }} – Arbeitnehmer –')
    clauses = [
        ('§ 1 Wiederaufnahme', 'Das ursprünglich am {{ original_contract_date }} begründete und durch Aufhebungsvertrag vom {{ termination_agreement_date }} zum {{ previous_end_date }} beendete Arbeitsverhältnis wird mit Wirkung zum {{ restart_date }} erneut aufgenommen.'),
        ('§ 2 Bedingungen', 'Das Arbeitsverhältnis wird zu den vereinbarten Bedingungen wiederaufgenommen. Beschäftigungsart: {{ employment_type }}. Befreiung von der Rentenversicherungspflicht beantragt: {{ pension_exemption_requested }}.'),
        ('§ 3 Sozialversicherung / Steuer', 'Der Arbeitnehmer bestätigt die Richtigkeit seiner Angaben zu weiteren Beschäftigungen und verpflichtet sich, Änderungen unverzüglich mitzuteilen.'),
    ]
    for h, t in clauses:
        p = d.add_paragraph(); p.add_run(h).bold = True; d.add_paragraph(t)
    add_signature(d)
    return save_doc(d, 'Wiederaufnahme - Muster .docx')


def paragraph_pdf(filename: str, title: str, blocks: list[str]) -> Path:
    path = WORK / filename
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm)
    story = [Paragraph(title, styles['Title']), Spacer(1, 10)]
    for block in blocks:
        story.extend([Paragraph(block, styles['BodyText']), Spacer(1, 8)])
    doc.build(story)
    return path


def form_pdf(filename: str, minijob=False) -> Path:
    path = WORK / filename
    c = canvas.Canvas(str(path), pagesize=A4)
    form: AcroForm = c.acroform
    width, height = A4
    y = height - 24*mm
    c.setFont('Helvetica-Bold', 14)
    c.drawString(18*mm, y, 'Personalfragebogen Minijobber (Stand 06.2025)' if minijob else 'Personalfragebogen (Stand 06.2025)')
    y -= 11*mm
    c.setFont('Helvetica', 8)
    c.drawString(18*mm, y, 'Die Angaben werden für die korrekte lohnsteuer- und sozialversicherungsrechtliche Beurteilung benötigt.')
    y -= 12*mm

    def text(label, name, x=18*mm, w=75*mm):
        nonlocal y
        c.setFont('Helvetica', 7); c.drawString(x, y+4*mm, label)
        form.textfield(name=name, x=x, y=y-1*mm, width=w, height=6*mm, borderWidth=0.5, fontSize=8)

    rows = [
        (('Vorname','Vorname'), ('Nachname','Nachname')),
        (('E-Mail-Adresse','E-Mail-Adresse'), ('Telefonnummer','Telefonnummer')),
        (('Straße und Hausnummer','Straße und Hausnummer'), ('Postleitzahl','Postleitzahl')),
        (('Ort','Ort'), ('Geburtsdatum','Geburtsdatum')),
        (('Staatsangehörigkeit','Staatsangehörigkeit'), ('Geburtsname','Geburtsname')),
        (('Geburtsort','Geburtsort'), ('Geburtsland','Geburtsland')),
        (('Sozial-/Rentenversicherungsnummer','Sozial-/Rentenversicherungsnummer'), ('Krankenkasse','Krankenkasse')),
        (('Steueridentifikationsnummer','Steueridentifikationsnummer'), ('IBAN','IBAN')),
        (('Kontoinhaber','Kontoinhaber'), ('Name des Kreditinstitus','Name des Kreditinstitus')),
    ]
    for left, right in rows:
        text(*left); text(*right, x=108*mm); y -= 12*mm
    c.setFont('Helvetica', 8); c.drawString(18*mm, y+3*mm, 'Anrede:')
    for idx, label in enumerate(['Herr','Frau','Divers','Unbestimmt']):
        form.radio(name='Optionsfeld 1', value=str(idx), selected=False, x=(40+idx*31)*mm, y=y, buttonStyle='circle', size=4*mm)
        c.drawString((45+idx*31)*mm, y+1*mm, label)
    y -= 10*mm
    c.drawString(18*mm, y+3*mm, 'Versicherung:')
    for idx, label in enumerate(['gesetzlich','privat']):
        form.radio(name='Optionsfeld 2', value=str(idx), selected=False, x=(50+idx*45)*mm, y=y, buttonStyle='circle', size=4*mm)
        c.drawString((55+idx*45)*mm, y+1*mm, label)
    if not minijob:
        y -= 10*mm; c.drawString(18*mm, y+3*mm, 'Kinder vorhanden:')
        for idx, label in enumerate(['ja','nein']):
            form.radio(name='Optionsfeld 3', value=str(idx), selected=False, x=(55+idx*40)*mm, y=y, buttonStyle='circle', size=4*mm)
            c.drawString((60+idx*40)*mm, y+1*mm, label)
        text('Falls ja: Wie viele Kinder?', 'Falls ja: Wie viele Kinder?', x=120*mm, w=65*mm)
    y -= 13*mm
    text('Ort, Datum', 'Ort, Datum 1. Seite', w=95*mm)
    c.line(120*mm, y+1*mm, 190*mm, y+1*mm); c.drawString(120*mm, y+4*mm, 'Unterschrift Mitarbeiter')
    c.showPage()
    c.setFont('Helvetica-Bold', 12); c.drawString(18*mm, height-22*mm, 'Weitere Beschäftigungsverhältnisse')
    y = height-38*mm
    for i in (1,2):
        text(f'{i}. Arbeitgeber weitere Tätigkeit', f'Name des Unternehmens/Arbeitgebers der weiteren Tätigkeit {i}', w=100*mm)
        text('Beginn', f'Beginn der Beschäftigung {i}', x=125*mm, w=60*mm); y -= 12*mm
        text('Verdienst brutto/Monat', f'Verdienst in Euro pro Monat {i}', w=80*mm); y -= 17*mm
    if minijob:
        c.setFont('Helvetica-Bold', 12); c.drawString(18*mm, y, 'Antrag auf Befreiung von der Rentenversicherungspflicht')
        y -= 13*mm
        text('Name Arbeitnehmer', 'Name des Arbeitnehmers', w=80*mm)
        text('Vorname Arbeitnehmer', 'Vorname des Arbeitnehmers', x=108*mm, w=77*mm); y -= 12*mm
        text('Rentenversicherungsnummer', 'Rentenversicherungsnummer', w=100*mm); y -= 15*mm
        c.setFont('Helvetica', 8)
        paragraph = ('Hiermit beantrage ich die Befreiung von der Versicherungspflicht in der Rentenversicherung im Rahmen meiner geringfügig entlohnten Beschäftigung. '
                     'Die Hinweise zu den möglichen Folgen der Befreiung habe ich zur Kenntnis genommen. Mir ist bekannt, dass der Antrag für gleichzeitig ausgeübte Minijobs bindend ist.')
        for line in [paragraph[i:i+115] for i in range(0, len(paragraph), 115)]:
            c.drawString(18*mm, y, line); y -= 4*mm
        y -= 5*mm
        text('Ort, Datum Arbeitnehmer', 'Ort, Datum Arbeitnehmer', w=90*mm); y -= 12*mm
        text('Name Arbeitgeber', 'Name des Arbeitgebers', w=90*mm)
        text('Betriebsnummer', 'Betriebsnummer', x=115*mm, w=70*mm); y -= 12*mm
        text('Eingang Befreiungsantrag', 'Befreiungsantrag Eingang', w=80*mm)
        text('Befreiung wirksam ab', 'Befreiung wirkt ab', x=108*mm, w=77*mm); y -= 12*mm
        text('Ort, Datum Arbeitgeber', 'Ort, Datum Arbeitgeber', w=90*mm)
    c.save()
    return path


def make_data_secrecy() -> Path:
    blocks = [
        '<b>Verpflichtung auf das Datengeheimnis (Stand: 22.05.2018)</b>',
        'Neben den betrieblichen Verschwiegenheitspflichten gelten die Europäische Datenschutz-Grundverordnung (DSGVO) und das Bundesdatenschutzgesetz (BDSG). Personenbezogene Daten dürfen nur verarbeitet werden, wenn eine Einwilligung oder gesetzliche Grundlage dies erlaubt.',
        'Personenbezogene Daten müssen rechtmäßig, transparent, zweckgebunden, datenminimiert, sachlich richtig, nur so lange wie erforderlich gespeichert und durch angemessene technische und organisatorische Maßnahmen geschützt werden.',
        'Verstöße können nach Art. 83 DSGVO in Verbindung mit §§ 41–43 BDSG geahndet werden und arbeitsrechtliche sowie zivilrechtliche Folgen haben. Die Verpflichtung besteht auch nach einem Aufgabenwechsel oder dem Ausscheiden fort.',
        'Mitarbeiter: {{ employee_name }}',
        'Ort, Datum: {{ signature_place }}, {{ signature_date }}<br/><br/>Unterschrift Mitarbeiter: ______________________________',
    ]
    return paragraph_pdf('1. Datengeheimnis.pdf', 'Verpflichtung auf das Datengeheimnis', blocks)


def download_leaflet() -> Path:
    path = WORK / '3. Merkblatt AN.pdf'
    url = 'https://www.arbeitsagentur.de/datei/merkblatt-leiharbeit_ba033905.pdf'
    req = urllib.request.Request(url, headers={'User-Agent':'APlusSolution/1.0'})
    with urllib.request.urlopen(req, timeout=60) as response:
        content = response.read()
    if not content.startswith(b'%PDF'):
        raise RuntimeError('Das offizielle BA-Merkblatt konnte nicht als PDF geladen werden.')
    path.write_bytes(content)
    return path


files = {
    'personalfragebogen-standard': form_pdf('Personalfragebogen A+ .pdf', False),
    'personalfragebogen-minijob': form_pdf('Personalfragebogen MJ Muster .pdf', True),
    'arbeitsvertrag-dgb-gvp': make_employment_contract(),
    'aufhebungsvertrag': make_termination(),
    'merkblatt-leiharbeitnehmer': download_leaflet(),
    'verpflichtung-datengeheimnis': make_data_secrecy(),
    'zusatzvereinbarung-entgelt': make_addendum(),
    'wiederaufnahme-arbeitsverhaeltnis': make_rehire(),
}
versions = {
    'personalfragebogen-standard':'06.2025',
    'personalfragebogen-minijob':'06.2025',
    'arbeitsvertrag-dgb-gvp':'11.2025',
    'aufhebungsvertrag':'1.0',
    'merkblatt-leiharbeitnehmer':'07.2026',
    'verpflichtung-datengeheimnis':'22.05.2018',
    'zusatzvereinbarung-entgelt':'1.0',
    'wiederaufnahme-arbeitsverhaeltnis':'1.0',
}
manifest = {'templates': []}
for slug, path in files.items():
    manifest['templates'].append({'slug': slug, 'file': path.name, 'version': versions[slug], 'sha256': sha(path)})
with zipfile.ZipFile(OUT, 'w', zipfile.ZIP_DEFLATED) as archive:
    archive.writestr('manifest.json', json.dumps(manifest, ensure_ascii=False, indent=2))
    for path in files.values():
        archive.write(path, path.name)
print(json.dumps({'bundle': str(OUT), 'templates': len(files), 'sha256': sha(OUT)}, ensure_ascii=False))
