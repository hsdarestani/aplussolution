import hashlib
import io
import json
from pathlib import Path

import requests
from django.conf import settings
from django.db import transaction
from django.utils import timezone
from docx import Document as DocxDocument
from pypdf import PdfReader

from .models import ClientOrder, Contract, Shift, ShiftImportPackage, ShiftImportRevision
from .native_cutover import approve_order
from .native_workforce import (
    _old_package_shifts_are_replaceable,
    _resolve_location,
    _resolve_position,
    package_shifts,
)
from .order_automation import (
    _parse_local_datetime,
    _validate_parsed,
    extract_request_id,
    fallback_request_id,
    normalize_request_id,
    payload_hash,
    resolve_client,
)

MAX_ORDER_FILE_BYTES = 15 * 1024 * 1024
MAX_ORDER_TEXT_CHARS = 180_000
ALLOWED_ORDER_EXTENSIONS = {'.pdf', '.docx', '.txt'}


def _clean_page(text):
    return '\n'.join(line.rstrip() for line in str(text or '').splitlines()).strip()


def extract_order_document(upload):
    """Extract machine-readable text from a staffing/order document.

    The source bytes are only used for parsing and are not persisted. PDF pages are
    kept separate so the AI result can be tied back to the exact source page during
    the review step.
    """
    if not upload:
        raise ValueError('Bitte zuerst eine Datei auswählen.')
    name = str(getattr(upload, 'name', '') or 'auftrag').strip()
    extension = Path(name).suffix.lower()
    if extension not in ALLOWED_ORDER_EXTENSIONS:
        raise ValueError('Erlaubt sind PDF, DOCX und TXT.')
    size = int(getattr(upload, 'size', 0) or 0)
    if size and size > MAX_ORDER_FILE_BYTES:
        raise ValueError('Die Datei ist größer als 15 MB.')

    data = upload.read()
    if len(data) > MAX_ORDER_FILE_BYTES:
        raise ValueError('Die Datei ist größer als 15 MB.')
    if not data:
        raise ValueError('Die Datei ist leer.')

    try:
        if extension == '.pdf':
            reader = PdfReader(io.BytesIO(data))
            pages = [_clean_page(page.extract_text() or '') for page in reader.pages]
        elif extension == '.docx':
            document = DocxDocument(io.BytesIO(data))
            rows = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells]
                    if any(values):
                        rows.append(' | '.join(values))
            pages = [_clean_page('\n'.join(rows))]
        else:
            try:
                text = data.decode('utf-8-sig')
            except UnicodeDecodeError:
                text = data.decode('latin-1')
            pages = [_clean_page(text)]
    except Exception as exc:
        raise ValueError('Die Datei konnte nicht gelesen werden.') from exc

    pages = [page for page in pages if page]
    if not pages:
        raise ValueError(
            'In der Datei wurde kein auslesbarer Text gefunden. Bei einem Scan bitte zuerst eine OCR-Version verwenden.'
        )
    text = '\n\n'.join(f'[[SEITE {index}]]\n{page}' for index, page in enumerate(pages, start=1))
    if len(text) > MAX_ORDER_TEXT_CHARS:
        raise ValueError('Das Dokument ist zu groß für einen einzelnen AI-Import.')
    return {'name': name, 'pages': pages, 'text': text}


def _document_prompt():
    return """Du bist die Auftrags- und Dienstplan-Automation von A+ Solution GmbH.
Analysiere ein deutsches Personal-Bestell-/Veranstaltungsdokument mit mehreren Seiten. Jede Veranstaltung bzw. jede Veranstaltungs-Nr. ist EIN eigener Auftrag. Vermische niemals zwei Veranstaltungsnummern.

Antworte ausschließlich als JSON-Objekt in dieser Form:
{
  "orders": [
    {
      "source_page": 1,
      "contract_no": "11059",
      "source_status": "Definitiv",
      "title": "Summerschool",
      "organizer": "...",
      "shifts": [
        {
          "role": "Servicekraft",
          "date": "2026-09-01",
          "start_time": "08:00",
          "end_time": "15:00",
          "count": 1,
          "location_text": "Evangelische Akademie Frankfurt",
          "site_text": "Evangelische Akademie Frankfurt",
          "site_address": "Römerberg 9, 60311 Frankfurt am Main",
          "notes": "Akademie"
        }
      ]
    }
  ]
}

Regeln:
- Extrahiere die Veranstaltungs-Nr. exakt; erfinde keine Nummer.
- Übernimm den Quellstatus exakt als "Definitiv" oder "Option", wenn vorhanden.
- Jede Zeile unter "Personal" wird als Schicht erfasst. Mehrere Personalzeilen ergeben mehrere Schichten.
- Eine Angabe wie "2 Servicekraft ... 14:00 - 21:00 (14 Std.)" bedeutet count=2; die Klammer enthält Gesamtstunden und NICHT die Anzahl.
- Bei mehrtägigen Veranstaltungen verwende für jede Personalzeile das Datum, unter dessen Tagesüberschrift sie steht.
- role soll die Personalrolle enthalten (z. B. Servicekraft, Serviceleitung, Garderobenkraft). Zusätze wie "Akademie", "vor Ort", besondere Kundenwünsche oder namentliche Wünsche gehören in notes und dürfen nicht verloren gehen.
- location_text/site_text sollen den tatsächlichen Einsatzort der Veranstaltung enthalten, nicht bloß die Rechnungsadresse des Veranstalters.
- Wenn auf einer Seite kein eigener Veranstaltungsort steht, nutze den bestmöglichen im Dokument genannten Einsatzort und erfinde keine Adresse.
- Erfasse ALLE Seiten und ALLE Personalzeilen. Lass keine Veranstaltung stillschweigend weg.
"""


def parse_order_document(document, session=None):
    if not settings.WIW_OPENAI_KEY:
        raise ValueError('OPENAI_API_KEY ist nicht konfiguriert.')
    client = session or requests.Session()
    response = client.post(
        'https://api.openai.com/v1/chat/completions',
        headers={
            'Authorization': f'Bearer {settings.WIW_OPENAI_KEY}',
            'Content-Type': 'application/json',
        },
        json={
            'model': settings.WIW_OPENAI_MODEL,
            'response_format': {'type': 'json_object'},
            'messages': [
                {'role': 'system', 'content': _document_prompt()},
                {'role': 'user', 'content': document['text']},
            ],
            'temperature': 0.05,
        },
        timeout=max(int(settings.WIW_HTTP_TIMEOUT), 90),
    )
    if not response.ok:
        raise ValueError(f'AI-Dokumentanalyse fehlgeschlagen ({response.status_code}).')
    body = response.json()
    content = (((body.get('choices') or [{}])[0].get('message') or {}).get('content'))
    try:
        parsed = json.loads(content)
    except (TypeError, ValueError) as exc:
        raise ValueError('AI hat kein gültiges JSON für das Dokument zurückgegeben.') from exc

    raw_orders = parsed.get('orders') if isinstance(parsed, dict) else None
    if not isinstance(raw_orders, list) or not raw_orders:
        raise ValueError('Im Dokument wurden keine Aufträge erkannt.')

    orders = []
    seen = set()
    pages = document['pages']
    for index, raw_order in enumerate(raw_orders, start=1):
        if not isinstance(raw_order, dict):
            continue
        try:
            page_no = int(raw_order.get('source_page') or index)
        except (TypeError, ValueError):
            page_no = index
        page_no = max(1, min(page_no, len(pages)))
        raw_text = pages[page_no - 1]
        clean = _validate_parsed({
            'contract_no': raw_order.get('contract_no') or '',
            'shifts': raw_order.get('shifts') or [],
        })
        request_id = extract_request_id(raw_text, clean) or normalize_request_id(raw_order.get('contract_no'))
        if not request_id:
            request_id = fallback_request_id(clean, raw_order.get('organizer') or raw_order.get('title') or '')
        # A duplicate event number means the model split one source order twice. Merge
        # the shifts instead of presenting two packages with the same immutable key.
        if request_id in seen:
            existing = next(item for item in orders if item['request_id'] == request_id)
            existing['shifts'].extend(clean['shifts'])
            continue
        seen.add(request_id)
        status = str(raw_order.get('source_status') or '').strip()
        if status.lower() not in {'definitiv', 'option'}:
            status = status or 'Unklar'
        orders.append({
            **clean,
            'request_id': request_id,
            'contract_no': request_id,
            'source_status': status,
            'source_page': page_no,
            'title': str(raw_order.get('title') or '').strip(),
            'organizer': str(raw_order.get('organizer') or '').strip(),
            'raw_text': raw_text,
        })

    if not orders:
        raise ValueError('Im Dokument wurden keine verwertbaren Aufträge erkannt.')
    orders.sort(key=lambda item: (item['shifts'][0]['date'], item['request_id']))
    return {
        'file_name': document['name'],
        'page_count': len(pages),
        'order_count': len(orders),
        'shift_count': sum(len(item['shifts']) for item in orders),
        'staff_slots': sum(sum(int(shift.get('count') or 1) for shift in item['shifts']) for item in orders),
        'orders': orders,
    }


def _status_tagged_order(order):
    status = str(order.get('source_status') or 'Unklar').strip() or 'Unklar'
    clean = _validate_parsed({'contract_no': order.get('contract_no') or order.get('request_id') or '', 'shifts': order.get('shifts') or []})
    for shift in clean['shifts']:
        note = str(shift.get('notes') or '').strip()
        marker = f'Quellstatus: {status}'
        if marker.lower() not in note.lower():
            shift['notes'] = (note + '\n' + marker).strip()
    clean['contract_no'] = str(order.get('contract_no') or order.get('request_id') or '').strip()
    return clean


def _approve_draft_order(parsed, raw_text, actor=None, client_id=None):
    """Store an unconfirmed source order as draft shifts without worker notifications."""
    parsed = _validate_parsed(parsed)
    client_hint = str(client_id or parsed['shifts'][0].get('site_text') or '')
    request_id = extract_request_id(raw_text, parsed) or fallback_request_id(parsed, client_hint)
    base_hash = payload_hash(parsed, client_hint, request_id)
    source_hash = hashlib.sha256(f'{base_hash}|draft'.encode()).hexdigest()

    with transaction.atomic():
        existing = ShiftImportPackage.objects.select_for_update().select_related('contract').filter(request_id=request_id).first()
        if existing and existing.source_hash == source_hash:
            return {'status': 'unchanged', 'source': 'aplus', 'mode': 'draft', 'request_id': request_id, 'package_id': str(existing.id)}
        if existing and existing.contract_id and existing.contract.status in {Contract.Status.SENT, Contract.Status.SIGNED}:
            raise ValueError('Der Auftrag ist bereits mit einem versendeten oder unterzeichneten Vertrag verknüpft.')
        if existing:
            replaceable, reason = _old_package_shifts_are_replaceable(existing)
            if not replaceable:
                raise ValueError(f'Der Optionsauftrag kann nicht ersetzt werden: {reason}')

        client = resolve_client(parsed, client_id)
        prepared = []
        for item in parsed['shifts']:
            start = _parse_local_datetime(item['date'], item['start_time'])
            end = _parse_local_datetime(item['date'], item['end_time'])
            if end <= start:
                from datetime import timedelta
                end += timedelta(days=1)
            prepared.append({
                **item,
                'start_dt': start,
                'end_dt': end,
                'location': _resolve_location(client, item),
                'position': _resolve_position(item.get('role')),
            })

        first_start = min(item['start_dt'] for item in prepared)
        last_end = max(item['end_dt'] for item in prepared)
        requested_staff = sum(max(1, int(item.get('count') or 1)) for item in prepared)
        functions = list(dict.fromkeys(item['position'].name for item in prepared))

        order = None
        old_payload = dict(existing.payload or {}) if existing else {}
        if existing:
            order_id = old_payload.get('order_id')
            if order_id:
                order = ClientOrder.objects.filter(pk=order_id).first()
        if not order:
            order = ClientOrder()
        order.client = client
        order.title = f'Auftrag {request_id} (Option)'
        order.description = raw_text
        order.location = prepared[0]['location']
        order.starts_at = first_start
        order.ends_at = last_end
        order.requested_staff = requested_staff
        order.functions = functions
        order.status = ClientOrder.Status.PLANNING
        if not order.created_by_id:
            order.created_by = actor
        order.save()

        old_ids = []
        if existing:
            old_ids = [str(item.id) for item in package_shifts(existing)]
            package_shifts(existing).delete()

        created_shifts = []
        for item in prepared:
            notes = (str(item.get('notes') or '').strip() + f'\nAuftrag: {request_id}\nManaged by A+ Workforce').strip()
            shift = Shift.objects.create(
                order=order,
                client=client,
                location=item['location'],
                position=item['position'],
                starts_at=item['start_dt'],
                ends_at=item['end_dt'],
                break_minutes=0,
                status=Shift.Status.DRAFT,
                is_open=True,
                notes=notes,
                required_count=max(1, int(item.get('count') or 1)),
                published_at=None,
            )
            created_shifts.append({
                'shift_id': str(shift.id),
                'local_shift_id': str(shift.id),
                'date': item['start_dt'].date().isoformat(),
                'start_time': item['start_dt'].strftime('%H:%M'),
                'end_time': item['end_dt'].strftime('%H:%M'),
                'role': item['position'].name,
                'location_id': str(item['location'].id),
                'location_name': item['location'].name,
                'position_id': str(item['position'].id),
                'required_count': shift.required_count,
                'notes': str(item.get('notes') or ''),
            })

        saved_payload = {
            'source_system': 'aplus',
            'source_status': 'Option',
            'request_id': request_id,
            'contract_no': extract_request_id(raw_text, parsed),
            'order_id': str(order.id),
            'site_name': client.name,
            'site_address': client.address or prepared[0]['location'].address,
            'client_type': 'company',
            'first_shift_time': first_start.isoformat(),
            'first_shift_end_time': last_end.isoformat(),
            'source_hash': source_hash,
            'shifts': created_shifts,
        }
        package, created = ShiftImportPackage.objects.update_or_create(
            request_id=request_id,
            defaults={
                'client': client,
                'site_name': client.name,
                'site_address': client.address or prepared[0]['location'].address,
                'first_shift_time': first_start,
                'first_shift_end_time': last_end,
                'raw_text': raw_text,
                'source_hash': source_hash,
                'payload': saved_payload,
                'status': ShiftImportPackage.Status.PENDING,
                'created_by': actor,
            },
        )
        version = (package.revisions.order_by('-version').values_list('version', flat=True).first() or 0) + 1
        ShiftImportRevision.objects.create(
            package=package,
            version=version,
            action='created' if created else 'updated',
            old_shift_ids=old_ids,
            new_shift_ids=[item['local_shift_id'] for item in created_shifts],
            old_payload=old_payload,
            new_payload=saved_payload,
        )

    return {
        'status': 'ok',
        'source': 'aplus',
        'mode': 'draft',
        'action': 'created' if created else 'updated',
        'version': version,
        'request_id': request_id,
        'package_id': str(package.id),
        'order_id': str(order.id),
        'shift_count': len(created_shifts),
        'created_count': requested_staff,
    }


def approve_document_orders(orders, actor=None, client_id=None):
    results = []
    errors = []
    for index, order in enumerate(orders or [], start=1):
        request_id = str(order.get('request_id') or order.get('contract_no') or f'#{index}')
        raw_text = str(order.get('raw_text') or f'Veranstaltungs-Nr.: {request_id}')
        source_status = str(order.get('source_status') or '').strip()
        try:
            clean = _status_tagged_order(order)
            if source_status.lower() == 'definitiv':
                result = approve_order(clean, raw_text, actor=actor, client_id=client_id)
                result['mode'] = 'published'
            else:
                result = _approve_draft_order(clean, raw_text, actor=actor, client_id=client_id)
            result['source_status'] = source_status or 'Unklar'
            results.append(result)
        except Exception as exc:
            errors.append({'request_id': request_id, 'detail': str(exc)})
    return {
        'status': 'ok' if not errors else ('partial' if results else 'failed'),
        'imported_orders': len(results),
        'published_orders': sum(1 for item in results if item.get('mode') == 'published'),
        'draft_orders': sum(1 for item in results if item.get('mode') == 'draft'),
        'created_staff_slots': sum(int(item.get('created_count') or 0) for item in results),
        'results': results,
        'errors': errors,
    }
