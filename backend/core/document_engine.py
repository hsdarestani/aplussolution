import hashlib
import io
import json
import os
import shutil
import subprocess
import tempfile
import zipfile
from datetime import date, datetime
from decimal import Decimal
from pathlib import Path

from django.conf import settings
from django.core.files.base import ContentFile
from django.template import Context, Template
from django.utils import timezone
from docx import Document as DocxDocument
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .document_catalog import CATALOG_BY_SLUG, DOCUMENT_CATALOG
from .models import Contract, ContractSignature, ContractTemplate, EmployeeMasterData


class DocumentGenerationError(ValueError):
    pass


PDF_FIELD_MAP = {
    'personalfragebogen-standard': {
        'first_name': 'Vorname', 'last_name': 'Nachname', 'email': 'E-Mail-Adresse',
        'phone': 'Telefonnummer', 'street': 'Straße und Hausnummer', 'postal_code': 'Postleitzahl',
        'city': 'Ort', 'birth_date': 'Geburtsdatum', 'nationality': 'Staatsangehörigkeit',
        'birth_name': 'Geburtsname', 'birth_place': 'Geburtsort', 'birth_country': 'Geburtsland',
        'social_insurance_number': 'Sozial-/Rentenversicherungsnummer',
        'health_insurance_name': 'Krankenkasse', 'tax_identification_number': 'Steueridentifikationsnummer',
        'tax_class': 'Steuerklasse/Faktor', 'tax_allowance': 'Freibeträge',
        'child_allowance': 'Kinderfreibetrag', 'religion': 'Konfession',
        'spouse_religion': 'Konfession des Ehepartners', 'iban': 'IBAN',
        'bank_account_holder': 'Kontoinhaber', 'bank_name': 'Name des Kreditinstitus',
        'children_under_25': 'Falls ja: Wie viele Kinder?', 'date_place': 'Ort, Datum 1. Seite',
    },
    'personalfragebogen-minijob': {
        'first_name': 'Vorname', 'last_name': 'Nachname', 'email': 'E-Mail-Adresse',
        'phone': 'Telefonnummer', 'street': 'Straße und Hausnummer', 'postal_code': 'Postleitzahl',
        'city': 'Ort', 'birth_date': 'Geburtsdatum', 'nationality': 'Staatsangehörigkeit',
        'birth_name': 'Geburtsname', 'birth_place': 'Geburtsort', 'birth_country': 'Geburtsland',
        'social_insurance_number': 'Sozial-/Rentenversicherungsnummer',
        'health_insurance_name': 'Krankenkasse', 'tax_identification_number': 'Steueridentifikationsnummer',
        'iban': 'IBAN', 'bank_account_holder': 'Kontoinhaber', 'bank_name': 'Name des Kreditinstitus',
        'date_place': 'Ort, Datum 1. Seite', 'employee_last_name': 'Name des Arbeitnehmers',
        'employee_first_name': 'Vorname des Arbeitnehmers', 'pension_number': 'Rentenversicherungsnummer',
        'employee_date_place': 'Ort, Datum Arbeitnehmer', 'employer_date_place': 'Ort, Datum Arbeitgeber',
        'employer_company_name': 'Name des Arbeitgebers', 'employer_business_number': 'Betriebsnummer',
        'request_received_on': 'Befreiungsantrag Eingang', 'exemption_effective_on': 'Befreiung wirkt ab',
    },
}


CHECKBOX_MAP = {
    'personalfragebogen-standard': {
        'salutation': ('Optionsfeld 1', {'herr': '/0', 'frau': '/1', 'divers': '/2', 'unbestimmt': '/3'}),
        'insurance_type': ('Optionsfeld 2', {'gesetzlich': '/0', 'privat': '/1'}),
        'has_children': ('Optionsfeld 3', {True: '/0', False: '/1', 'true': '/0', 'false': '/1'}),
    },
    'personalfragebogen-minijob': {
        'salutation': ('Optionsfeld 1', {'herr': '/0', 'frau': '/1', 'divers': '/2', 'unbestimmt': '/3'}),
        'insurance_type': ('Optionsfeld 2', {'gesetzlich': '/0', 'privat': '/1'}),
    },
}


def format_value(value):
    if value in (None, ''):
        return ''
    if isinstance(value, bool):
        return 'Ja' if value else 'Nein'
    if isinstance(value, (date, datetime)):
        return value.strftime('%d.%m.%Y')
    if isinstance(value, Decimal):
        return f'{value:.2f}'.replace('.', ',')
    if isinstance(value, list):
        return '; '.join(str(item) for item in value)
    return str(value)


def deep_value(data, path, default=''):
    value = data
    for part in path.split('.'):
        if not isinstance(value, dict):
            return default
        value = value.get(part)
    return default if value is None else value


def company_data():
    return {
        'name': settings.COMPANY_NAME,
        'address': settings.COMPANY_ADDRESS,
        'business_number': getattr(settings, 'COMPANY_BUSINESS_NUMBER', ''),
        'aueg_license_authority': settings.AUEG_LICENSE_AUTHORITY,
        'aueg_license_date': settings.AUEG_LICENSE_DATE,
    }


def contract_data(contract):
    worker = contract.worker
    user = worker.user if worker else None
    master_obj = EmployeeMasterData.objects.filter(worker=worker).first() if worker else None
    master = dict(master_obj.data or {}) if master_obj else {}
    master.setdefault('full_address', ' '.join(str(item) for item in [master.get('street'), master.get('postal_code'), master.get('city')] if item).strip())
    base = {
        'today': timezone.localdate(),
        'company': company_data(),
        'contract': {
            'id': str(contract.id),
            'title': contract.title,
            'starts_on': contract.starts_on,
            'ends_on': contract.ends_on,
            'variables': contract.variables or {},
        },
        'user': {
            'first_name': user.first_name if user else '',
            'last_name': user.last_name if user else '',
            'email': user.email if user else '',
            'phone': user.phone if user else '',
        },
        'worker': {
            'full_name': user.get_full_name() if user else '',
            'employee_number': worker.employee_number if worker else '',
            'employment_type': worker.employment_type if worker else '',
            'monthly_hours': worker.monthly_hours if worker else '',
            'tariff_hourly_rate': worker.tariff_hourly_rate if worker else '',
            'extra_allowance': worker.extra_allowance if worker else '',
        },
        'master': master,
        'client': {
            'name': contract.client.name if contract.client else '',
            'address': contract.client.address if contract.client else '',
            'customer_number': contract.client.customer_number if contract.client else '',
            'vat_id': contract.client.vat_id if contract.client else '',
        },
    }
    output = {}
    for item in contract.template.schema.get('fields', []):
        source = item.get('source')
        value = deep_value(base, source) if source else ''
        if item['name'] in (contract.variables or {}):
            value = contract.variables[item['name']]
        output[item['name']] = value
    output.update(contract.variables or {})
    employment_type = output.get('employment_type') or base['worker']['employment_type']
    output.update({
        'employment_full_151': employment_type == 'vollzeit' and str(output.get('monthly_hours') or '') in {'151.67', '151,67'},
        'employment_full_173': employment_type == 'vollzeit' and str(output.get('monthly_hours') or '') in {'173.34', '173,34'},
        'employment_part': employment_type in {'teilzeit', 'student'},
        'employment_minijob': employment_type == 'minijob',
        'company_name': base['company']['name'],
        'company_address': base['company']['address'],
        'employee_name': base['worker']['full_name'],
        'employee_address': base['master'].get('full_address', ''),
        'first_name': base['user']['first_name'],
        'last_name': base['user']['last_name'],
        'email': base['user']['email'],
        'phone': base['user']['phone'],
        'date': base['today'],
        'signature_date': base['today'],
    })
    signatures = {item.role: item for item in contract.signatures.all()}
    for role in ContractSignature.Role.values:
        signature = signatures.get(role)
        output[f'{role}_signature_name'] = signature.signer_name if signature else ''
        output[f'{role}_signature_data'] = signature.signature_data if signature else ''
        output[f'{role}_signed_at'] = signature.signed_at if signature else ''
    return output


def validate_required_fields(template, data):
    missing = []
    for item in template.schema.get('fields', []):
        if item.get('required') and data.get(item['name']) in (None, '', [], {}):
            missing.append({'field': item['name'], 'label': item.get('label', item['name'])})
    if missing:
        labels = ', '.join(item['label'] for item in missing)
        raise DocumentGenerationError(f'Pflichtangaben fehlen: {labels}')
    return missing


def seed_document_catalog():
    result = {'created': 0, 'updated': 0}
    for item in DOCUMENT_CATALOG:
        defaults = {
            'name': item['name'],
            'kind': item['kind'],
            'audience': item['audience'],
            'version': item['version'],
            'schema': {'fields': item.get('fields', []), 'signature_roles': item.get('signature_roles', []), 'source_name': item.get('source_name'), 'overlay': item.get('overlay', {})},
            'source_format': item['source_format'],
            'requires_signature': item['requires_signature'],
            'required_document': True,
            'active': True,
        }
        obj, created = ContractTemplate.objects.update_or_create(slug=item['slug'], defaults=defaults)
        result['created' if created else 'updated'] += 1
    return result


def import_template_bundle(file_obj):
    created, updated, errors = 0, 0, []
    with zipfile.ZipFile(file_obj) as archive:
        if 'manifest.json' not in archive.namelist():
            raise DocumentGenerationError('manifest.json fehlt im Vorlagenpaket.')
        manifest = json.loads(archive.read('manifest.json').decode('utf-8-sig'))
        for entry in manifest.get('templates', []):
            try:
                slug = entry['slug']
                catalog = CATALOG_BY_SLUG.get(slug)
                if not catalog:
                    raise DocumentGenerationError(f'Unbekannte Vorlage: {slug}')
                filename = entry['file']
                content = archive.read(filename)
                checksum = hashlib.sha256(content).hexdigest()
                if entry.get('sha256') and not hmac_safe_equal(checksum, entry['sha256']):
                    raise DocumentGenerationError(f'Prüfsumme stimmt nicht: {filename}')
                obj = ContractTemplate.objects.get(slug=slug)
                obj.source_file.save(Path(filename).name, ContentFile(content), save=False)
                obj.source_checksum = checksum
                obj.version = entry.get('version') or obj.version
                obj.save()
                updated += 1
            except Exception as exc:
                errors.append({'entry': entry.get('slug') or entry.get('file'), 'error': str(exc)})
    return {'created': created, 'updated': updated, 'errors': errors}


def hmac_safe_equal(left, right):
    import hmac
    return hmac.compare_digest(str(left).lower(), str(right).lower())


def replace_text_in_paragraph(paragraph, values):
    original = ''.join(run.text for run in paragraph.runs)
    updated = original
    for key, value in values.items():
        updated = updated.replace('{{ ' + key + ' }}', format_value(value))
        updated = updated.replace('{{' + key + '}}', format_value(value))
        updated = updated.replace('{{checkbox:' + key + '}}', '☒' if bool(value) else '☐')
    if updated == original:
        return
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(updated)


def render_docx(source_path, data):
    document = DocxDocument(source_path)
    for paragraph in document.paragraphs:
        replace_text_in_paragraph(paragraph, data)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, data)
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


def convert_docx_to_pdf(docx_bytes):
    binary = getattr(settings, 'LIBREOFFICE_BINARY', 'libreoffice')
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = Path(tmp) / 'document.docx'
        docx_path.write_bytes(docx_bytes)
        result = subprocess.run(
            [binary, '--headless', '--convert-to', 'pdf', '--outdir', tmp, str(docx_path)],
            capture_output=True,
            text=True,
            timeout=120,
        )
        pdf_path = Path(tmp) / 'document.pdf'
        if result.returncode != 0 or not pdf_path.exists():
            raise DocumentGenerationError(f'DOCX konnte nicht in PDF umgewandelt werden: {result.stderr[-300:]}')
        return pdf_path.read_bytes()


def fill_pdf_form(source_path, data, slug):
    reader = PdfReader(source_path)
    writer = PdfWriter(clone_from=reader)
    mapping = PDF_FIELD_MAP.get(slug, {})
    fields = {}
    for key, field_name in mapping.items():
        if key == 'date_place':
            place = data.get('signature_place') or data.get('city') or ''
            value = f'{format_value(place)}, {format_value(data.get("date"))}'.strip(', ')
        elif key == 'employee_date_place':
            value = f'{format_value(data.get("signature_place") or data.get("city"))}, {format_value(data.get("date"))}'.strip(', ')
        elif key == 'employer_date_place':
            value = f'{format_value(data.get("signature_place") or "Frankfurt am Main")}, {format_value(data.get("date"))}'.strip(', ')
        elif key == 'employee_last_name':
            value = data.get('last_name')
        elif key == 'employee_first_name':
            value = data.get('first_name')
        elif key == 'pension_number':
            value = data.get('social_insurance_number')
        else:
            value = data.get(key)
        fields[field_name] = format_value(value)
    for key, (field_name, options) in CHECKBOX_MAP.get(slug, {}).items():
        raw = data.get(key)
        normalized = raw.lower() if isinstance(raw, str) else raw
        if normalized in options:
            fields[field_name] = options[normalized]
    other = data.get('other_employments') or []
    for index, employment in enumerate(other[:2], start=1):
        if not isinstance(employment, dict):
            continue
        fields[f'Name des Unternehmens/Arbeitgebers der weiteren Tätigkeit {index}'] = format_value(employment.get('employer'))
        fields[f'Beginn der Beschäftigung {index}'] = format_value(employment.get('start_date'))
        fields[f'Verdienst in Euro pro Monat {index}'] = format_value(employment.get('gross_monthly'))
    writer.set_need_appearances_writer()
    for page in writer.pages:
        writer.update_page_form_field_values(page, fields, auto_regenerate=True)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def overlay_footer(source_path, data):
    reader = PdfReader(source_path)
    writer = PdfWriter()
    for index, page in enumerate(reader.pages):
        if index == len(reader.pages) - 1:
            packet = io.BytesIO()
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            c = canvas.Canvas(packet, pagesize=(width, height))
            c.setFont('Helvetica', 9)
            c.drawString(55, 45, f"{format_value(data.get('signature_place'))}, {format_value(data.get('signature_date'))}")
            if data.get('employee_signature_name'):
                c.drawString(300, 45, format_value(data.get('employee_signature_name')))
            c.save()
            packet.seek(0)
            overlay = PdfReader(packet).pages[0]
            page.merge_page(overlay)
        writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def render_html_pdf(contract, data):
    html = Template(contract.template.html_template).render(Context({**data, 'contract': contract}))
    plain = html.replace('<h1>', '').replace('</h1>', '\n\n').replace('<h2>', '<b>').replace('</h2>', '</b>\n').replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=20 * mm, leftMargin=20 * mm, topMargin=18 * mm, bottomMargin=18 * mm, title=contract.title, author=settings.COMPANY_NAME)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='ContractTitle', parent=styles['Title'], alignment=TA_CENTER, spaceAfter=14))
    story = [Paragraph(contract.title, styles['ContractTitle'])]
    for block in plain.split('\n\n'):
        if block.strip():
            story.extend([Paragraph(block.strip().replace('\n', '<br/>'), styles['BodyText']), Spacer(1, 7)])
    doc.build(story)
    return buffer.getvalue()


def generate_contract_files(contract, validate=True):
    data = contract_data(contract)
    if validate:
        validate_required_fields(contract.template, data)
    if not contract.template.source_file and contract.template.source_format != ContractTemplate.SourceFormat.HTML:
        raise DocumentGenerationError(f'Quelldatei für „{contract.template.name}“ ist noch nicht installiert.')
    source_path = contract.template.source_file.path if contract.template.source_file else None
    docx_bytes = None
    if contract.template.source_format == ContractTemplate.SourceFormat.DOCX:
        docx_bytes = render_docx(source_path, data)
        pdf_bytes = convert_docx_to_pdf(docx_bytes)
    elif contract.template.source_format == ContractTemplate.SourceFormat.PDF_OVERLAY:
        slug = contract.template.slug
        if slug in PDF_FIELD_MAP:
            pdf_bytes = fill_pdf_form(source_path, data, slug)
        else:
            pdf_bytes = overlay_footer(source_path, data)
    elif contract.template.source_format == ContractTemplate.SourceFormat.STATIC_PDF:
        pdf_bytes = Path(source_path).read_bytes()
    else:
        pdf_bytes = render_html_pdf(contract, data)
    previous_status = contract.status
    if docx_bytes:
        contract.docx.save(f'{contract.template.slug}-{contract.id}.docx', ContentFile(docx_bytes), save=False)
    contract.pdf.save(f'{contract.template.slug}-{contract.id}.pdf', ContentFile(pdf_bytes), save=False)
    contract.data_snapshot = {key: format_value(value) for key, value in data.items()}
    contract.generated_at = timezone.now()
    if previous_status == Contract.Status.SIGNED:
        contract.status = Contract.Status.SIGNED
    elif previous_status == Contract.Status.SENT:
        contract.status = Contract.Status.SENT
    else:
        contract.status = Contract.Status.READY
    contract.save()
    return contract


def generate_worker_packet(worker, created_by=None, variables=None):
    seed_document_catalog()
    contracts = []
    for template in ContractTemplate.objects.filter(active=True, required_document=True, audience__in=[ContractTemplate.Audience.WORKER, ContractTemplate.Audience.BOTH]).order_by('name'):
        contract = Contract.objects.create(
            template=template,
            worker=worker,
            title=template.name,
            variables=variables or {},
            source_system='wiw' if worker.wiw_user_id else 'aplus',
            created_by=created_by,
        )
        try:
            generate_contract_files(contract)
        except DocumentGenerationError:
            pass
        contracts.append(contract)
    return contracts
