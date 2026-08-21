import io
from datetime import date

import pytest
from docx import Document

from core import document_engine
from core.smart_docx import render_docx_bytes


def _source_docx():
    document = Document()

    paragraph = document.add_paragraph()
    paragraph.add_run('Mitarbeiter: ')
    start = paragraph.add_run('{{employee')
    start.bold = True
    end = paragraph.add_run('_name}}')
    end.bold = True
    suffix = paragraph.add_run(' · Beginn: ')
    suffix.italic = True
    paragraph.add_run('{{ start_date }}')

    multiline = document.add_paragraph('Vereinbarung: ')
    multiline.add_run('{{additional_agreements}}')

    checkbox = document.add_paragraph('Bestätigt: ')
    checkbox.add_run('{{check')
    checkbox.add_run('box:confirmed}}')

    table = document.add_table(rows=1, cols=1)
    cell_paragraph = table.cell(0, 0).paragraphs[0]
    cell_paragraph.add_run('Adresse: {{employee_')
    cell_paragraph.add_run('address}}')

    header = document.sections[0].header.paragraphs[0]
    header.add_run('Firma: {{company_')
    header.add_run('name}}')

    footer = document.sections[0].footer.paragraphs[0]
    footer.add_run('Dokument vom {{date}}')

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _formatter(value):
    if isinstance(value, date):
        return value.strftime('%d.%m.%Y')
    return '' if value is None else str(value)


def test_smart_docx_replaces_split_runs_and_preserves_surrounding_formatting():
    rendered, diagnostics = render_docx_bytes(
        _source_docx(),
        {
            'employee_name': 'Max Mustermann',
            'start_date': date(2026, 9, 1),
            'additional_agreements': 'Zeile eins\nZeile zwei',
            'confirmed': 'false',
            'employee_address': 'Mainzer Landstraße 123, 60329 Frankfurt',
            'company_name': 'A+ Solution GmbH',
            'date': date(2026, 8, 21),
        },
        _formatter,
    )

    result = Document(io.BytesIO(rendered))
    first = result.paragraphs[0]
    assert first.text == 'Mitarbeiter: Max Mustermann · Beginn: 01.09.2026'
    assert 'Max Mustermann' in first.runs[1].text
    assert first.runs[1].bold is True
    assert any(run.italic and 'Beginn:' in run.text for run in first.runs)

    assert result.paragraphs[1].text == 'Vereinbarung: Zeile eins\nZeile zwei'
    assert result.paragraphs[2].text == 'Bestätigt: ☐'
    assert result.tables[0].cell(0, 0).text == 'Adresse: Mainzer Landstraße 123, 60329 Frankfurt'
    assert result.sections[0].header.paragraphs[0].text == 'Firma: A+ Solution GmbH'
    assert result.sections[0].footer.paragraphs[0].text == 'Dokument vom 21.08.2026'
    assert diagnostics['replacements'] == 7
    assert diagnostics['unresolved_keys'] == []
    assert 'word/document.xml' in diagnostics['processed_parts']
    assert any(name.startswith('word/header') for name in diagnostics['processed_parts'])
    assert any(name.startswith('word/footer') for name in diagnostics['processed_parts'])


def test_smart_docx_reports_unknown_placeholders_without_destroying_them():
    document = Document()
    document.add_paragraph('Bekannt {{known}} · Unbekannt {{missing_key}}')
    source = io.BytesIO()
    document.save(source)

    rendered, diagnostics = render_docx_bytes(source.getvalue(), {'known': 'OK'}, str)
    result = Document(io.BytesIO(rendered))

    assert result.paragraphs[0].text == 'Bekannt OK · Unbekannt {{missing_key}}'
    assert diagnostics['unresolved_keys'] == ['missing_key']


def test_contract_document_engine_uses_smart_docx_renderer_and_rejects_unknown_tags(tmp_path):
    assert getattr(document_engine.render_docx, '_smart_docx_renderer', False) is True

    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run('{{employee')
    paragraph.add_run('_name}}')
    source = tmp_path / 'contract.docx'
    document.save(source)

    rendered = document_engine.render_docx(source, {'employee_name': 'Erika Musterfrau'})
    assert Document(io.BytesIO(rendered)).paragraphs[0].text == 'Erika Musterfrau'

    bad_document = Document()
    bad_document.add_paragraph('{{unknown_contract_field}}')
    bad_source = tmp_path / 'bad-contract.docx'
    bad_document.save(bad_source)

    with pytest.raises(document_engine.DocumentGenerationError, match='unknown_contract_field'):
        document_engine.render_docx(bad_source, {})
