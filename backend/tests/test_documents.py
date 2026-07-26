import hashlib
import io
import json
import zipfile
from datetime import date, timedelta
from pathlib import Path
from unittest.mock import patch

import pytest
from django.core import mail
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.utils import timezone
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas

from core.document_engine import (
    DocumentGenerationError,
    contract_data,
    generate_contract_files,
    import_template_bundle,
    render_docx,
    seed_document_catalog,
    validate_required_fields,
)
from core.models import Contract, ContractSignature, ContractTemplate, EmployeeMasterData, Notification
from core.services import sign_contract
from core.tasks import send_contract_reminders


def make_docx_bytes(text='Hallo {{ employee_name }} {{checkbox:neuanstellung}}'):
    doc = DocxDocument()
    doc.add_paragraph(text)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def make_pdf_bytes():
    output = io.BytesIO()
    c = canvas.Canvas(output)
    c.drawString(50, 800, 'Static PDF')
    c.save()
    return output.getvalue()


@pytest.mark.django_db
def test_catalog_seeds_exactly_eight_documents():
    result = seed_document_catalog()
    assert ContractTemplate.objects.count() == 8
    assert result['created'] == 8
    second = seed_document_catalog()
    assert second['updated'] == 8
    assert ContractTemplate.objects.filter(slug='arbeitsvertrag-dgb-gvp').exists()


@pytest.mark.django_db
def test_private_bundle_checksum_validation(tmp_path):
    seed_document_catalog()
    content = make_docx_bytes()
    checksum = hashlib.sha256(content).hexdigest()
    bundle = io.BytesIO()
    with zipfile.ZipFile(bundle, 'w') as archive:
        archive.writestr('manifest.json', json.dumps({'templates': [{'slug': 'arbeitsvertrag-dgb-gvp', 'file': 'template.docx', 'sha256': checksum, 'version': '11.2025'}]}))
        archive.writestr('template.docx', content)
    bundle.seek(0)
    result = import_template_bundle(bundle)
    assert result['updated'] == 1 and not result['errors']
    template = ContractTemplate.objects.get(slug='arbeitsvertrag-dgb-gvp')
    assert template.source_checksum == checksum
    bad = io.BytesIO()
    with zipfile.ZipFile(bad, 'w') as archive:
        archive.writestr('manifest.json', json.dumps({'templates': [{'slug': 'arbeitsvertrag-dgb-gvp', 'file': 'template.docx', 'sha256': 'bad'}]}))
        archive.writestr('template.docx', content)
    bad.seek(0)
    result = import_template_bundle(bad)
    assert result['errors']


@pytest.mark.django_db
def test_docx_token_replacement_preserves_document_structure(tmp_path):
    source = tmp_path / 'source.docx'
    source.write_bytes(make_docx_bytes())
    output = render_docx(source, {'employee_name': 'Anna Becker', 'neuanstellung': True})
    rendered = DocxDocument(io.BytesIO(output))
    assert 'Anna Becker' in rendered.paragraphs[0].text
    assert '☒' in rendered.paragraphs[0].text


@pytest.mark.django_db
def test_required_fields_report_labels(worker_user, admin_user):
    template = ContractTemplate.objects.create(name='Required', slug='required', kind='employment', schema={'fields': [{'name': 'iban', 'label': 'IBAN', 'required': True, 'source': 'master.iban'}]}, html_template='x')
    contract = Contract.objects.create(template=template, worker=worker_user.worker_profile, title='Test', created_by=admin_user)
    data = contract_data(contract)
    with pytest.raises(DocumentGenerationError, match='IBAN'):
        validate_required_fields(template, data)


@pytest.mark.django_db
def test_static_pdf_generation(worker_user, admin_user):
    template = ContractTemplate.objects.create(name='Leaflet', slug='leaflet-test', kind='aueg_leaflet', source_format='static_pdf', schema={'fields': [], 'signature_roles': []}, requires_signature=False)
    template.source_file.save('leaflet.pdf', ContentFile(make_pdf_bytes()))
    contract = Contract.objects.create(template=template, worker=worker_user.worker_profile, title='Merkblatt', created_by=admin_user)
    generate_contract_files(contract)
    contract.refresh_from_db()
    assert contract.pdf and contract.generated_at
    assert contract.status == 'ready'


@pytest.mark.django_db
def test_two_party_signature_completes_only_after_both(api_client, worker_user, admin_user):
    template = ContractTemplate.objects.create(name='Agreement', slug='agreement-test', kind='termination', source_format='html', html_template='<h1>{{ employee_name }}</h1>', schema={'fields': [], 'signature_roles': ['employee', 'employer']})
    contract = Contract.objects.create(template=template, worker=worker_user.worker_profile, title='Agreement', status='ready', created_by=admin_user)
    api_client.force_authenticate(worker_user)
    first = api_client.post(f'/api/contracts/{contract.id}/sign/', {'name': 'Anna Becker', 'signature': 'Anna'}, format='json')
    assert first.status_code == 200
    contract.refresh_from_db()
    assert contract.status == 'sent'
    api_client.force_authenticate(admin_user)
    second = api_client.post(f'/api/contracts/{contract.id}/sign/', {'name': 'A+ Solution GmbH', 'signature': 'Admin'}, format='json')
    assert second.status_code == 200
    contract.refresh_from_db()
    assert contract.status == 'signed'
    assert set(contract.signatures.values_list('role', flat=True)) == {'employee', 'employer'}


@pytest.mark.django_db
def test_contract_send_notifies_worker_client_and_admin(auth_admin, admin_user, worker_user, client_user, company):
    template = ContractTemplate.objects.create(name='HTML', slug='notify-template', kind='employment', source_format='html', html_template='<h1>Test</h1>', schema={'fields': [], 'signature_roles': ['employee', 'employer']})
    contract = Contract.objects.create(template=template, worker=worker_user.worker_profile, client=company, title='Notify', created_by=admin_user)
    response = auth_admin.post(f'/api/contracts/{contract.id}/send/', {}, format='json')
    assert response.status_code == 200
    for user in [admin_user, worker_user, client_user]:
        assert Notification.objects.filter(user=user, kind=f'contract-sent-{contract.id}').exists()


@pytest.mark.django_db
def test_contract_reminder_sends_to_all_recipients(admin_user, worker_user, client_user, company):
    template = ContractTemplate.objects.create(name='Reminder', slug='reminder-template', kind='employment', html_template='x', schema={})
    contract = Contract.objects.create(template=template, worker=worker_user.worker_profile, client=company, title='Ending', status='sent', ends_on=timezone.localdate()+timedelta(days=7), created_by=admin_user)
    sent = send_contract_reminders()
    assert sent >= 3
    assert len(mail.outbox) >= 1
    assert Notification.objects.filter(kind=f'contract-7-{contract.id}').count() >= 3


@pytest.mark.django_db
def test_generate_worker_packet_creates_eight_records(auth_admin, worker_user):
    seed_document_catalog()
    response = auth_admin.post(f'/api/workers/{worker_user.worker_profile.id}/generate_packet/', {'variables': {}}, format='json')
    assert response.status_code == 201
    assert Contract.objects.filter(worker=worker_user.worker_profile).count() == 8
